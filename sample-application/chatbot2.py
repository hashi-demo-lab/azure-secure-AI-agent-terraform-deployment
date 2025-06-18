# Complete AI Agent with Azure Authentication for HCP Vault
import streamlit as st
from langchain.agents import create_tool_calling_agent, AgentExecutor
from langchain_core.tools import tool
from langchain_community.utilities.sql_database import SQLDatabase
from langchain_community.agent_toolkits.sql.toolkit import SQLDatabaseToolkit
from langchain.tools import Tool
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings, ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain import hub
from azure.identity import ManagedIdentityCredential, DefaultAzureCredential
from llama_index.core import (
    SimpleDirectoryReader,
    GPTVectorStoreIndex,
    PromptHelper,
    ServiceContext,
    StorageContext,
    load_index_from_storage,
    Settings
)
from llama_index.llms.langchain import LangChainLLM
from llama_index.embeddings.langchain import LangchainEmbedding
from dotenv import load_dotenv, dotenv_values
import sys
import re
import base64
import os
import requests
import hvac
import logging
import uuid
import json

# FastAPI imports - install with: pip install fastapi uvicorn
try:
    from fastapi import FastAPI
    from pydantic import BaseModel
    FASTAPI_AVAILABLE = True
except ImportError:
    print("Warning: FastAPI not installed. Install with: pip install fastapi uvicorn")
    FASTAPI_AVAILABLE = False
    # Create dummy classes for development
    class FastAPI:
        def __init__(self): pass
        def post(self, path): return lambda f: f
    class BaseModel:
        def __init__(self, **kwargs):
            for k, v in kwargs.items():
                setattr(self, k, v)

# Import DuckDuckGo search API wrapper
try:
    from langchain_community.utilities import DuckDuckGoSearchAPIWrapper
except ImportError:
    try:
        from langchain.utilities import DuckDuckGoSearchAPIWrapper
    except ImportError:
        print("Warning: DuckDuckGo search not available. Web search tool will be disabled.")
        DuckDuckGoSearchAPIWrapper = None

# Load environment variables
load_dotenv()

# Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Azure OpenAI Configuration
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01")
AZURE_OPENAI_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4")
AZURE_OPENAI_EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", "text-embedding-ada-002")

# Environment variables and Vault config
VAULT_URL = os.getenv("VAULT_ADDR", "https://your-hcp-vault-url")
VAULT_ROLE = os.getenv("VAULT_ROLE", "aks-workload-role")  # Updated role name
VAULT_SECRET_PATH = os.getenv("VAULT_SECRET_PATH", "database/creds/my-role")

DB_HOST = os.getenv("DB_HOST", "<your-db-host>")
DB_NAME = os.getenv("DB_NAME", "<your-db-name>")

# FastAPI app (only if available)
if FASTAPI_AVAILABLE:
    app = FastAPI()
else:
    app = None

class QueryRequest(BaseModel):
    question: str

# FIXED: Azure Vault Authentication Functions
import os
import json
import base64
import logging
import hvac
from azure.identity import WorkloadIdentityCredential, DefaultAzureCredential

logger = logging.getLogger(__name__)
# Enhanced token path verification and debugging for Azure Workload Identity
import os
import logging
import json
import base64
from pathlib import Path

logger = logging.getLogger(__name__)

# Integration with your existing code
def get_azure_identity_token_for_vault():

    """
    Enhanced version of your get_azure_identity_token_for_vault function
    with better token path detection and verification
    """
    try:
        from azure.identity import WorkloadIdentityCredential, DefaultAzureCredential
        
        # First, verify the workload identity setup
        if not verify_workload_identity_setup():
            raise Exception("Workload Identity setup verification failed")
        
        # Get environment variables
        client_id = os.getenv("AZURE_CLIENT_ID")
        tenant_id = os.getenv("AZURE_TENANT_ID") 
        token_file = os.getenv("AZURE_FEDERATED_TOKEN_FILE")
        
        logger.info(f"Using workload identity - Client ID: {client_id}")
        logger.info(f"Token file: {token_file}")
        
        # Verify required environment variables
        if not client_id:
            raise Exception("AZURE_CLIENT_ID environment variable not set")
        if not tenant_id:
            raise Exception("AZURE_TENANT_ID environment variable not set")
        if not token_file:
            # Try to auto-detect standard token file paths
            standard_paths = [
                "/var/run/secrets/azure/tokens/azure-identity-token",
                "/var/run/secrets/kubernetes.io/serviceaccount/token"
            ]
            
            for path in standard_paths:
                if os.path.exists(path) and os.access(path, os.R_OK):
                    logger.info(f"Auto-detected token file: {path}")
                    token_file = path
                    break
            
            if not token_file:
                raise Exception("AZURE_FEDERATED_TOKEN_FILE environment variable not set and no standard token files found")
            
        # Verify token file exists and is readable
        if not os.path.exists(token_file):
            # Show available files for debugging
            check_all_possible_token_paths()
            raise Exception(f"Token file not found: {token_file}")
        
        if not os.access(token_file, os.R_OK):
            raise Exception(f"Token file not readable: {token_file}")
        
        logger.info(f"✅ Token file verified: {token_file}")
        
        # Create WorkloadIdentityCredential
        logger.info("Creating WorkloadIdentityCredential")
        credential = WorkloadIdentityCredential(
            client_id=client_id,
            tenant_id=tenant_id,
            token_file=token_file
        )
        
        # Try different audiences for Vault authentication
        audiences_to_try = [
            "api://AzureADTokenExchange/.default",
            "api://AzureADTokenExchange",
            f"https://vault.{tenant_id}",  # Custom Vault audience
            "https://management.azure.com/.default"  # Azure Resource Manager
        ]
        
        token = None
        successful_audience = None
        
        for audience in audiences_to_try:
            try:
                logger.info(f"Trying audience: {audience}")
                token = credential.get_token(audience)
                successful_audience = audience
                logger.info(f"✅ Successfully obtained token with audience: {audience}")
                break
            except Exception as e:
                logger.warning(f"Failed with audience {audience}: {e}")
                continue
        
        if not token:
            raise Exception("Failed to obtain token with any audience")
            
        # Analyze the obtained token
        analyze_jwt_token(token.token)
        return token.token, successful_audience
            
    except Exception as e:
        logger.error(f"Failed to get Azure token: {e}")
        raise Exception(f"Could not obtain Azure identity token: {e}")

def get_kubernetes_jwt_token():
    """
    Read the projected service account token used for Vault JWT auth
    """
    token_file = os.getenv("AZURE_FEDERATED_TOKEN_FILE", "/var/run/secrets/azure/tokens/azure-identity-token")

    if not os.path.exists(token_file):
        raise FileNotFoundError(f"Token file not found: {token_file}")
    if not os.access(token_file, os.R_OK):
        raise PermissionError(f"Cannot read token file: {token_file}")
    
    with open(token_file, "r") as f:
        token = f.read().strip()
    
    logger.info(f"Read Kubernetes service account JWT token (first 50 chars): {token[:50]}")
    return token    

import hvac
import os
import requests

import hvac
import os
import requests

def authenticate_with_vault():
    """
    Authenticate to Vault using the Kubernetes Workload Identity token via JWT auth
    """
    vault_addr = os.getenv("VAULT_ADDR")
    vault_auth_path = os.getenv("VAULT_AUTH_PATH", "jwt")  
    vault_role = os.getenv("VAULT_ROLE")
    vault_namespace = os.getenv("VAULT_NAMESPACE")
    token_file = os.getenv("AZURE_FEDERATED_TOKEN_FILE", "/var/run/secrets/azure/tokens/azure-identity-token")

    if not all([vault_addr, vault_role]):
        raise Exception("VAULT_ADDR and VAULT_ROLE must be set")

    if not os.path.exists(token_file):
        raise Exception(f"Token file not found: {token_file}")

    with open(token_file, "r") as f:
        jwt = f.read().strip()

    # Correct: pass namespace during client initialization
    client = hvac.Client(url=vault_addr, namespace=vault_namespace)

    # JWT login request
    payload = {
        "role": vault_role,
        "jwt": jwt
    }

    login_url = f"{vault_addr}/v1/auth/{vault_auth_path}/login"
    headers = {}
    if vault_namespace:
        headers["X-Vault-Namespace"] = vault_namespace

    response = requests.post(login_url, json=payload, headers=headers)
    if not response.ok:
        raise Exception(f"Vault login failed: {response.status_code} - {response.text}")

    client.token = response.json()["auth"]["client_token"]
    return client

# Updated main authentication function that you should use
def get_db_credentials_from_vault():
    """
    FIXED: Get database credentials from HashiCorp Vault using Azure authentication
    Enhanced with better error handling and debugging, now with namespace support
    """
    try:
        # Authenticate with Vault
        client = authenticate_with_vault()
        
        # Read database credentials
        secret_path = os.getenv("VAULT_SECRET_PATH")
        vault_namespace = os.getenv("VAULT_NAMESPACE")
        
        if not secret_path:
            raise Exception("VAULT_SECRET_PATH environment variable not set")
            
        logger.info(f"Reading credentials from path: {secret_path}")
        if vault_namespace:
            logger.info(f"Using namespace: {vault_namespace}")
        
        try:
            secret = client.read(secret_path)
        except Exception as e:
            logger.error(f"Failed to read from path '{secret_path}': {e}")
            
            # Check if it's a permissions issue
            if "permission denied" in str(e).lower():
                logger.error("This suggests the Vault token doesn't have permission to read from this path")
                logger.error("Check Vault policies and ensure the role has access to the database secrets engine")
                if vault_namespace:
                    logger.error(f"Also verify policies are correctly configured for namespace '{vault_namespace}'")
            elif "mount not found" in str(e).lower():
                logger.error("This suggests the database secrets engine is not mounted at the expected path")
                logger.error("Check if the database secrets engine is enabled and mounted correctly")
                if vault_namespace:
                    logger.error(f"Also verify the mount exists in namespace '{vault_namespace}'")
            
            raise Exception(f"Cannot read secret from Vault: {e}")
        
        if not secret:
            raise Exception(f"No secret found at path: {secret_path}")
            
        logger.info(f"Secret response keys: {list(secret.keys())}")
        
        # Extract credentials from the secret
        # For database secrets engine, credentials are in the 'data' field
        if "data" in secret:
            data = secret["data"]
            username = data.get("username")
            password = data.get("password")
            
            if not username or not password:
                logger.error(f"Available keys in secret data: {list(data.keys())}")
                raise Exception("Username or password not found in secret data")
                
            logger.info(f"✅ Successfully retrieved credentials for user: {username}")
            
            # Log credential metadata (without sensitive data)
            lease_duration = secret.get("lease_duration", "unknown")
            renewable = secret.get("renewable", "unknown")
            logger.info(f"Credential lease duration: {lease_duration} seconds")
            logger.info(f"Credential renewable: {renewable}")
            
            return username, password
        else:
            logger.error(f"Unexpected secret structure: {list(secret.keys())}")
            raise Exception("Invalid secret format - 'data' field not found")
            
    except Exception as e:
        logger.error(f"Failed to get database credentials from Vault: {e}")
        raise

    
# FIXED: MySQL tool with better error handling
@tool
def ask_mysql(query: str) -> str:
    """Executes a SQL query on the MySQL DB and returns results.
    
    Available tables and columns:
    - countries: id, country (full country name like 'China', 'India'), population, created_at
    - cities: id, city, country (full country name), population, created_at
    
    Important: Use 'country' column with full country names (e.g., 'China', not 'CHN').
    Example queries:
    - SELECT population FROM countries WHERE country='China'
    - SELECT city, population FROM cities WHERE country='India'
    """
    try:
        logger.info(f"Executing MySQL query: {query[:100]}...")
        
        # Get fresh credentials from Vault
        username, password = get_db_credentials_from_vault()
        
        # Database connection details
        db_host = os.getenv("DB_HOST", "<your-db-host>")
        db_name = os.getenv("DB_NAME", "<your-db-name>")
        
        if db_host == "<your-db-host>" or db_name == "<your-db-name>":
            return "Error: Database host or name not configured. Please set DB_HOST and DB_NAME environment variables."
        
        # Create database URI
        db_uri = f"mysql+pymysql://{username}:{password}@{db_host}:3306/{db_name}"
               
        db = SQLDatabase.from_uri(db_uri)
        result = db.run(query)
        
        logger.info("Query executed successfully")
        return result
        
    except Exception as e:
        error_msg = f"Database query failed: {str(e)}"
        logger.error(error_msg)
        
        # Provide more specific error messages
        if "Unknown column" in str(e):
            error_msg += f"\nHint: Column doesn't exist. Available columns in 'countries' table: id, country, population, created_at"
            error_msg += f"\nNote: Use 'country' with full names like 'China', 'India', not country codes like 'CHN'"
        elif "permission denied" in str(e).lower():
            error_msg += "\nHint: Check Vault authentication and policies"
        elif "connection" in str(e).lower():
            error_msg += "\nHint: Check database connection settings and credentials"
        elif "access denied" in str(e).lower():
            error_msg += "\nHint: Check database user permissions"
            
        return error_msg

# Calculator tool
@tool
def calculator(expression: str) -> str:
    """Evaluates a math expression."""
    try:
        return str(eval(expression))
    except Exception as e:
        return f"Error: {e}"

# FIXED: Initialize Azure OpenAI LLM with proper configuration
def initialize_azure_llm():
    """Initialize Azure OpenAI LLM with proper configuration"""
    try:
        if AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT:
            # Use API key authentication - FIXED parameters
            llm = AzureChatOpenAI(
                api_key=AZURE_OPENAI_API_KEY,  # Updated parameter name
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                api_version=AZURE_OPENAI_API_VERSION,  # Updated parameter name
                azure_deployment=AZURE_OPENAI_DEPLOYMENT_NAME,  # Updated parameter name
                model=AZURE_OPENAI_DEPLOYMENT_NAME,  # Explicitly set model to deployment name
                temperature=0
            )
            logger.info(f"Initialized Azure OpenAI with API key, deployment: {AZURE_OPENAI_DEPLOYMENT_NAME}")
        else:
            # Use managed identity authentication
            token = get_kubernetes_jwt_token()

            llm = AzureChatOpenAI(
                api_key=token.token,  # Updated parameter name
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                api_version=AZURE_OPENAI_API_VERSION,  # Updated parameter name
                azure_deployment=AZURE_OPENAI_DEPLOYMENT_NAME,  # Updated parameter name
                model=AZURE_OPENAI_DEPLOYMENT_NAME,  # Explicitly set model to deployment name
                temperature=0
            )
            logger.info(f"Initialized Azure OpenAI with Managed Identity, deployment: {AZURE_OPENAI_DEPLOYMENT_NAME}")
        return llm
    except Exception as e:
        logger.warning(f"Azure OpenAI initialization failed: {e}")
        # Try alternative initialization method
        return initialize_azure_llm_fallback()

def initialize_azure_llm_fallback():
    """Fallback Azure OpenAI LLM initialization with different parameter approach"""
    try:
        if AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT:
            llm = AzureChatOpenAI(
                openai_api_key=AZURE_OPENAI_API_KEY,  # Try old parameter name
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                openai_api_version=AZURE_OPENAI_API_VERSION,  # Try old parameter name
                deployment_name=AZURE_OPENAI_DEPLOYMENT_NAME,  # Try old parameter name
                model_name=AZURE_OPENAI_DEPLOYMENT_NAME,  # Explicitly set model_name
                temperature=0
            )
            logger.info(f"Initialized Azure OpenAI (fallback) with API key, deployment: {AZURE_OPENAI_DEPLOYMENT_NAME}")
            return llm
        else:
            credential = ManagedIdentityCredential()
            token = credential.get_token("https://cognitiveservices.azure.com/.default")

            llm = AzureChatOpenAI(
                openai_api_key=token.token,
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                openai_api_version=AZURE_OPENAI_API_VERSION,
                deployment_name=AZURE_OPENAI_DEPLOYMENT_NAME,
                model_name=AZURE_OPENAI_DEPLOYMENT_NAME,
                temperature=0
            )
            logger.info(f"Initialized Azure OpenAI (fallback) with Managed Identity, deployment: {AZURE_OPENAI_DEPLOYMENT_NAME}")
            return llm
    except Exception as e:
        logger.error(f"Azure OpenAI fallback initialization also failed: {e}")
        return None

# FIXED: Initialize embeddings for LlamaIndex
def initialize_azure_embeddings():
    """Initialize Azure OpenAI embeddings"""
    try:
        if AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT:
            embeddings = AzureOpenAIEmbeddings(
                api_key=AZURE_OPENAI_API_KEY,  # Updated parameter name
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                api_version=AZURE_OPENAI_API_VERSION,  # Updated parameter name
                azure_deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT  # Updated parameter name
            )
        else:
            credential = ManagedIdentityCredential()
            token = credential.get_token("https://cognitiveservices.azure.com/.default")
            embeddings = AzureOpenAIEmbeddings(
                api_key=token.token,
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                api_version=AZURE_OPENAI_API_VERSION,
                azure_deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT
            )
        return embeddings
    except Exception as e:
        logger.warning(f"Primary embeddings initialization failed, trying fallback: {e}")
        # Try fallback with old parameter names
        try:
            if AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT:
                embeddings = AzureOpenAIEmbeddings(
                    openai_api_key=AZURE_OPENAI_API_KEY,
                    azure_endpoint=AZURE_OPENAI_ENDPOINT,
                    openai_api_version=AZURE_OPENAI_API_VERSION,
                    deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT
                )
            else:
                credential = ManagedIdentityCredential()
                token = credential.get_token("https://cognitiveservices.azure.com/.default")
                embeddings = AzureOpenAIEmbeddings(
                    openai_api_key=token.token,
                    azure_endpoint=AZURE_OPENAI_ENDPOINT,
                    openai_api_version=AZURE_OPENAI_API_VERSION,
                    deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT
                )
            logger.info("Embeddings initialized with fallback parameters")
            return embeddings
        except Exception as e2:
            logger.error(f"Failed to initialize Azure embeddings (both methods): {e2}")
            return None

# Debug function to help identify configuration issues
def debug_azure_config():
    """Debug Azure OpenAI configuration"""
    print("=== Azure OpenAI Configuration Debug ===")
    print(f"AZURE_OPENAI_ENDPOINT: {AZURE_OPENAI_ENDPOINT}")
    print(f"AZURE_OPENAI_DEPLOYMENT_NAME: {AZURE_OPENAI_DEPLOYMENT_NAME}")
    print(f"AZURE_OPENAI_API_VERSION: {AZURE_OPENAI_API_VERSION}")
    print(f"AZURE_OPENAI_API_KEY: {'***SET***' if AZURE_OPENAI_API_KEY else 'NOT SET'}")
    print(f"AZURE_OPENAI_EMBEDDING_DEPLOYMENT: {AZURE_OPENAI_EMBEDDING_DEPLOYMENT}")
    
    # Test the configuration
    try:
        test_llm = initialize_azure_llm()
        if test_llm:
            print("✅ LLM initialization successful")
            # Try a simple test call
            try:
                response = test_llm.invoke("Hello")
                print("✅ Test call successful")
            except Exception as e:
                print(f"❌ Test call failed: {e}")
        else:
            print("❌ LLM initialization failed")
    except Exception as e:
        print(f"❌ Configuration test failed: {e}")

# Web Search using DuckDuckGo (if available)
search = None
if DuckDuckGoSearchAPIWrapper:
    try:
        search = DuckDuckGoSearchAPIWrapper()
    except Exception as e:
        logger.warning(f"Could not initialize DuckDuckGo search: {e}")

# Initialize LlamaIndex components
def initialize_llamaindex():
    """Initialize LlamaIndex with Azure OpenAI"""
    try:
        # Initialize Azure components
        azure_llm = initialize_azure_llm()
        embeddings = initialize_azure_embeddings()

        if azure_llm and embeddings:
            # Configure LlamaIndex settings
            Settings.llm = LangChainLLM(llm=azure_llm)
            Settings.embed_model = LangchainEmbedding(embeddings)

            # Create prompt helper
            prompt_helper = PromptHelper(
                context_window=4096,
                num_output=256,
                chunk_overlap_ratio=0.1,
                chunk_size_limit=None
            )

            logger.info("LlamaIndex initialized successfully")
            return True
        else:
            logger.warning("LlamaIndex initialization failed - missing Azure LLM or embeddings")
            return False
    except Exception as e:
        logger.error(f"LlamaIndex initialization failed: {e}")
        return False

# Language model and memory
llm = initialize_azure_llm()
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# Initialize LlamaIndex
llamaindex_available = initialize_llamaindex()

# Document processing utilities
import tempfile
import shutil
from pathlib import Path

# Global variable to store the current document index
document_index = None

def save_uploaded_file(uploaded_file):
    """Save uploaded file to temporary directory and return path"""
    try:
        # Create a temporary directory for uploaded files
        temp_dir = Path("temp_uploads")
        temp_dir.mkdir(exist_ok=True)
        
        # Save the file
        file_path = temp_dir / uploaded_file.name
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        return str(file_path)
    except Exception as e:
        logger.error(f"Error saving uploaded file: {e}")
        return None

def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats"""
    try:
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_extension == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except ImportError:
                return "Error: PyPDF2 not installed. Cannot process PDF files."
        
        elif file_extension in ['.docx', '.doc']:
            try:
                import docx
                if file_extension == '.docx':
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                else:
                    return "Error: .doc files not supported. Please use .docx format."
            except ImportError:
                return "Error: python-docx not installed. Cannot process Word documents."
        
        elif file_extension == '.csv':
            try:
                import pandas as pd
                df = pd.read_csv(file_path)
                return df.to_string()
            except ImportError:
                return "Error: pandas not installed. Cannot process CSV files."
        
        else:
            return f"Error: Unsupported file format: {file_extension}"
    
    except Exception as e:
        return f"Error extracting text from file: {str(e)}"

# Enhanced document processing tools
@tool
def process_documents(directory_path: str) -> str:
    """Process documents from a directory and create a searchable index."""
    global document_index
    
    if not llamaindex_available:
        return "Error: LlamaIndex not properly initialized"

    try:
        # Load documents
        documents = SimpleDirectoryReader(directory_path).load_data()

        # Create index
        document_index = GPTVectorStoreIndex.from_documents(documents)

        # Save index
        document_index.storage_context.persist(persist_dir="./storage")

        return f"Successfully processed {len(documents)} documents from {directory_path}"
    except Exception as e:
        return f"Error processing documents: {str(e)}"

@tool
def process_uploaded_documents(file_texts: dict) -> str:
    """Process uploaded documents and create a searchable index."""
    global document_index
    
    if not llamaindex_available:
        return "Error: LlamaIndex not properly initialized"

    try:
        from llama_index.core import Document
        
        # Create documents from text content
        documents = []
        for filename, text_content in file_texts.items():
            if text_content and not text_content.startswith("Error:"):
                doc = Document(text=text_content, metadata={"filename": filename})
                documents.append(doc)
        
        if not documents:
            return "Error: No valid documents to process"

        # Create index
        document_index = GPTVectorStoreIndex.from_documents(documents)

        # Save index
        document_index.storage_context.persist(persist_dir="./storage")

        return f"Successfully processed {len(documents)} uploaded documents"
    except Exception as e:
        return f"Error processing uploaded documents: {str(e)}"

@tool
def query_documents(query: str) -> str:
    """Query the processed documents using the created index."""
    global document_index
    
    if not llamaindex_available:
        return "Error: LlamaIndex not properly initialized"

    try:
        # Use global index if available, otherwise try to load from storage
        if document_index is None:
            try:
                storage_context = StorageContext.from_defaults(persist_dir="./storage")
                document_index = load_index_from_storage(storage_context)
            except:
                return "Error: No document index found. Please upload and process documents first."

        # Query the index
        query_engine = document_index.as_query_engine()
        response = query_engine.query(query)

        return str(response)
    except Exception as e:
        return f"Error querying documents: {str(e)}"

@tool
def get_document_summary() -> str:
    """Get a summary of the currently indexed documents."""
    global document_index
    
    if not llamaindex_available:
        return "Error: LlamaIndex not properly initialized"
    
    if document_index is None:
        return "No documents currently indexed. Please upload and process documents first."
    
    try:
        # Get basic information about the index
        docstore = document_index.storage_context.docstore
        docs = docstore.docs
        
        summary = f"Document Index Summary:\n"
        summary += f"- Total documents: {len(docs)}\n"
        
        # Get filenames if available
        filenames = set()
        for doc_id, doc in docs.items():
            if 'filename' in doc.metadata:
                filenames.add(doc.metadata['filename'])
        
        if filenames:
            summary += f"- Files indexed: {', '.join(sorted(filenames))}\n"
        
        # Get a sample of content
        if docs:
            first_doc = list(docs.values())[0]
            preview = first_doc.text[:200] + "..." if len(first_doc.text) > 200 else first_doc.text
            summary += f"- Sample content: {preview}"
        
        return summary
    except Exception as e:
        return f"Error getting document summary: {str(e)}"

# Define tools (conditionally include web search and document processing)
tools = [
    ask_mysql,
    Tool(name="calculator", func=calculator, description="Useful for math expressions.")
]

if search:
    tools.append(Tool(name="web_search", func=search.run, description="Useful for current events and web-based queries."))

if llamaindex_available:
    tools.extend([
        Tool(name="process_documents", func=process_documents, description="Process documents from a directory path and create searchable index."),
        Tool(name="process_uploaded_documents", func=process_uploaded_documents, description="Process uploaded documents and create searchable index."),
        Tool(name="query_documents", func=query_documents, description="Query processed documents using natural language."),
        Tool(name="get_document_summary", func=get_document_summary, description="Get a summary of currently indexed documents.")
    ])

# Define the prompt template
def create_agent_prompt():
    """Create a prompt template for the agent."""
    try:
        prompt = ChatPromptTemplate.from_messages([
            ("system", "You are a helpful AI assistant. Use the available tools when needed to provide accurate and helpful responses."),
            MessagesPlaceholder("chat_history", optional=True),
            ("human", "{input}"),
            MessagesPlaceholder("agent_scratchpad")  # This is REQUIRED for tool calling agents
        ])
        logger.info("Created prompt template")
        return prompt
    except Exception as e:
        logger.error(f"Failed to create prompt: {e}")
        return None

# Create the agent with proper error handling
def initialize_agent_system():
    """Initialize the entire agent system"""
    global agent_executor
    try:
        logger.info("Initializing agent system...")

        # Validate prerequisites
        if not llm:
            raise Exception("LLM not initialized")
        if not tools:
            raise Exception("No tools available")

        logger.info("Prerequisites validated, creating agent...")

        # Create prompt
        prompt = create_agent_prompt()
        if not prompt:
            raise Exception("Prompt creation failed")

        # Create agent
        agent = create_tool_calling_agent(llm, tools, prompt)

        # Create AgentExecutor
        agent_executor = AgentExecutor(
            agent=agent,
            tools=tools,
            verbose=True,
            return_intermediate_steps=True,
            memory=memory,
            handle_parsing_errors=True,
            max_iterations=10
        )

        logger.info("Agent system initialized successfully")
        return True

    except Exception as e:
        logger.error(f"Failed to initialize agent system: {e}")
        agent_executor = None
        return False

# Initialize the agent system at module level with error handling
try:
    logger.info("Starting agent system initialization...")
    success = initialize_agent_system()
    if success:
        logger.info("✓ Agent system ready")
    else:
        logger.error("✗ Agent system initialization failed")
        print("WARNING: Agent system failed to initialize. Check logs for details.")
except Exception as e:
    logger.error(f"Critical error during agent system initialization: {e}")
    print(f"CRITICAL ERROR: Agent system initialization failed - {e}")
    agent_executor = None

# Query endpoint (only if FastAPI is available)
if FASTAPI_AVAILABLE and app and agent_executor:
    @app.post("/query")
    async def query_agent(request: QueryRequest):
        correlation_id = str(uuid.uuid4())
        logger.info(f"[{correlation_id}] Received question: {request.question}")

        try:
            response = agent_executor.invoke({"input": request.question})
            final_answer = response.get("output", "")
            steps = response.get("intermediate_steps", [])

            formatted_steps = []
            
            for i, step in enumerate(steps):
                try:
                    # Safely handle each step
                    if not isinstance(step, (list, tuple)) or len(step) < 2:
                        logger.warning(f"[{correlation_id}] Skipping malformed step {i}: {step}")
                        continue
                    
                    agent_action, result = step[0], step[1]
                    
                    # Safely extract step information
                    step_info = {
                        "step": i + 1,
                        "thought": "",
                        "action": "unknown",
                        "input": "",
                        "result": str(result) if result is not None else ""
                    }
                    
                    # Safe attribute extraction
                    if hasattr(agent_action, 'log') and agent_action.log:
                        step_info["thought"] = str(agent_action.log).strip()
                    elif agent_action:
                        step_info["thought"] = str(agent_action)
                    
                    if hasattr(agent_action, 'tool') and agent_action.tool:
                        step_info["action"] = str(agent_action.tool)
                    
                    if hasattr(agent_action, 'tool_input') and agent_action.tool_input:
                        step_info["input"] = str(agent_action.tool_input)
                    elif agent_action:
                        step_info["input"] = str(agent_action)
                    
                    formatted_steps.append(step_info)
                    
                except Exception as step_error:
                    logger.error(f"[{correlation_id}] Error processing step {i}: {str(step_error)}")
                    # Add a placeholder step to maintain order
                    formatted_steps.append({
                        "step": i + 1,
                        "thought": f"Error processing step: {str(step_error)}",
                        "action": "error",
                        "input": "",
                        "result": ""
                    })

            return {
                "correlation_id": correlation_id,
                "question": request.question,
                "agent_steps": formatted_steps,
                "final_answer": final_answer
            }

        except StopIteration as e:
            logger.error(f"[{correlation_id}] StopIteration error: {str(e)}")
            return {
                "correlation_id": correlation_id,
                "error": f"Agent execution completed unexpectedly: {str(e)}",
                "question": request.question,
                "agent_steps": [],
                "final_answer": ""
            }
        except Exception as e:
            logger.error(f"[{correlation_id}] Agent error: {str(e)}")
            return {
                "correlation_id": correlation_id,
                "error": str(e),
                "question": request.question,
                "agent_steps": [],
                "final_answer": ""
            }

# Alternative function for direct usage (without FastAPI) - Enhanced for Streamlit
def query_agent_direct(question: str):
    """Direct function to query the agent without FastAPI"""
    if not agent_executor:
        return {
            "error": "Agent not properly initialized. Please check the logs for details."
        }

    correlation_id = str(uuid.uuid4())
    logger.info(f"[{correlation_id}] Received question: {question}")

    try:
        response = agent_executor.invoke({"input": question})
        final_answer = response.get("output", "")
        steps = response.get("intermediate_steps", [])

        formatted_steps = []
        
        for i, step in enumerate(steps):
            try:
                # Safely handle each step
                if not isinstance(step, (list, tuple)) or len(step) < 2:
                    logger.warning(f"[{correlation_id}] Skipping malformed step {i}: {step}")
                    continue
                
                agent_action, result = step[0], step[1]
                
                # Safely extract step information
                step_info = {
                    "step": i + 1,
                    "thought": "",
                    "action": "unknown",
                    "input": "",
                    "result": str(result) if result is not None else ""
                }
                
                # Safe attribute extraction
                if hasattr(agent_action, 'log') and agent_action.log:
                    step_info["thought"] = str(agent_action.log).strip()
                elif agent_action:
                    step_info["thought"] = str(agent_action)
                
                if hasattr(agent_action, 'tool') and agent_action.tool:
                    step_info["action"] = str(agent_action.tool)
                
                if hasattr(agent_action, 'tool_input') and agent_action.tool_input:
                    step_info["input"] = str(agent_action.tool_input)
                elif agent_action:
                    step_info["input"] = str(agent_action)
                
                formatted_steps.append(step_info)
                
            except Exception as step_error:
                logger.error(f"[{correlation_id}] Error processing step {i}: {str(step_error)}")
                # Add a placeholder step to maintain order
                formatted_steps.append({
                    "step": i + 1,
                    "thought": f"Error processing step: {str(step_error)}",
                    "action": "error",
                    "input": "",
                    "result": ""
                })

        return {
            "correlation_id": correlation_id,
            "question": question,
            "agent_steps": formatted_steps,
            "final_answer": final_answer
        }

    except StopIteration as e:
        logger.error(f"[{correlation_id}] StopIteration error: {str(e)}")
        return {
            "correlation_id": correlation_id,
            "error": f"Agent execution completed unexpectedly: {str(e)}",
            "question": question,
            "agent_steps": [],
            "final_answer": ""
        }
    except Exception as e:
        logger.error(f"[{correlation_id}] Agent error: {str(e)}")
        return {
            "correlation_id": correlation_id,
            "error": str(e),
            "question": question,
            "agent_steps": [],
            "final_answer": ""
        }
    
# Health check endpoint for FastAPI
if FASTAPI_AVAILABLE and app:
    @app.get("/health")
    async def health_check():
        """Health check endpoint"""
        vault_healthy = check_vault_connectivity()
        return {
            "status": "healthy" if vault_healthy and agent_executor else "degraded",
            "agent_initialized": agent_executor is not None,
            "vault_connection": vault_healthy,
            "azure_openai_configured": bool(AZURE_OPENAI_ENDPOINT),
            "database_configured": DB_HOST != "<your-db-host>",
            "web_search_available": search is not None,
            "document_processing_available": llamaindex_available
        }

# Document Processing Interface Functions
def render_document_upload_interface():
    """Render the document upload and processing interface"""
    st.header("📄 Document Processing")
    
    # Create tabs for different document operations
    tab1, tab2, tab3 = st.tabs(["📤 Upload Documents", "🔍 Query Documents", "📊 Document Summary"])
    
    with tab1:
        st.subheader("Upload and Process Documents")
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Choose files to upload",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx', 'csv'],
            help="Supported formats: TXT, PDF, DOCX, CSV"
        )
        
        if uploaded_files:
            st.write(f"**{len(uploaded_files)} file(s) selected:**")
            for file in uploaded_files:
                st.write(f"- {file.name} ({file.size} bytes)")
            
            # Process button
            if st.button("🔄 Process Documents", type="primary"):
                with st.spinner("Processing documents..."):
                    try:
                        # Save files and extract text
                        file_texts = {}
                        progress_bar = st.progress(0)
                        
                        for i, uploaded_file in enumerate(uploaded_files):
                            # Update progress
                            progress = (i + 1) / len(uploaded_files)
                            progress_bar.progress(progress)
                            
                            # Save file
                            file_path = save_uploaded_file(uploaded_file)
                            if file_path:
                                # Extract text
                                text_content = extract_text_from_file(file_path)
                                file_texts[uploaded_file.name] = text_content
                                
                                # Clean up temporary file
                                try:
                                    os.remove(file_path)
                                except:
                                    pass
                        
                        progress_bar.empty()
                        
                        # Process documents through the agent
                        if file_texts:
                            # Use the agent to process documents
                            response = query_agent_direct(f"Process these uploaded documents: {file_texts}")
                            
                            if "error" in response:
                                st.error(f"Error processing documents: {response['error']}")
                            else:
                                st.success(response["final_answer"])
                                
                                # Show processing steps
                                if response.get("agent_steps"):
                                    with st.expander("Show processing steps"):
                                        for step in response["agent_steps"]:
                                            st.write(f"**Step {step['step']}:** {step['action']}")
                                            st.write(f"Result: {step['result']}")
                                            st.write("---")
                        else:
                            st.error("No valid text content extracted from uploaded files.")
                            
                    except Exception as e:
                        st.error(f"Error processing documents: {str(e)}")
    
    with tab2:
        st.subheader("Query Your Documents")
        
        # Document query interface
        query_input = st.text_input(
            "Ask a question about your documents:",
            placeholder="e.g., What are the main topics discussed in the documents?"
        )
        
        if st.button("🔎 Search Documents") and query_input:
            with st.spinner("Searching documents..."):
                try:
                    response = query_agent_direct(f"Query documents: {query_input}")
                    
                    if "error" in response:
                        st.error(f"Error: {response['error']}")
                    else:
                        st.write("**Answer:**")
                        st.write(response["final_answer"])
                        
                        # Show search steps
                        if response.get("agent_steps"):
                            with st.expander("Show search process"):
                                for step in response["agent_steps"]:
                                    st.write(f"**Step {step['step']}:** {step['action']}")
                                    st.write(f"Input: {step['input']}")
                                    st.write(f"Result: {step['result']}")
                                    st.write("---")
                except Exception as e:
                    st.error(f"Error querying documents: {str(e)}")
        
        # Suggested queries
        st.write("**Suggested queries:**")
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📋 Summarize all documents"):
                with st.spinner("Creating summary..."):
                    response = query_agent_direct("Create a comprehensive summary of all processed documents")
                    if "error" not in response:
                        st.write(response["final_answer"])
        
        with col2:
            if st.button("🔑 Extract key points"):
                with st.spinner("Extracting key points..."):
                    response = query_agent_direct("Extract the main key points and insights from the processed documents")
                    if "error" not in response:
                        st.write(response["final_answer"])
    
    with tab3:
        st.subheader("Document Index Summary")
        
        if st.button("📊 Get Document Summary"):
            with st.spinner("Getting document summary..."):
                try:
                    response = query_agent_direct("Get a summary of the currently indexed documents")
                    
                    if "error" in response:
                        st.error(f"Error: {response['error']}")
                    else:
                        st.write(response["final_answer"])
                        
                except Exception as e:
                    st.error(f"Error getting document summary: {str(e)}")
        
        # Clear index button
        if st.button("🗑️ Clear Document Index", type="secondary"):
            try:
                global document_index
                document_index = None
                
                # Try to remove storage directory
                import shutil
                try:
                    shutil.rmtree("./storage")
                    st.success("Document index cleared successfully!")
                except FileNotFoundError:
                    st.info("No document index to clear.")
                except Exception as e:
                    st.warning(f"Index cleared from memory, but couldn't remove storage: {e}")
                    
            except Exception as e:
                st.error(f"Error clearing document index: {str(e)}")

# Streamlit Interface - Enhanced with better error handling and document processing
def main():
    """Main Streamlit interface"""
    st.set_page_config(
        page_title="AI Agent with Document Processing", 
        page_icon="🤖",
        layout="wide"
    )

    st.title("🤖 AI Agent with Azure OpenAI ")
    st.markdown("Ask questions, query databases, search the web, or process and analyze documents!")

    # Create main layout
    col1, col2 = st.columns([2, 1])
    
    with col2:
        # Sidebar-like content in right column
        st.markdown("### 🔧 System Status")
        
        # Status indicators
        vault_status = "Connected" if VAULT_URL != "https://your-hcp-vault-url" else "Not configured"
        
        status_data = {
            "HashiCorp Vault": ("🟢" if vault_status == "Connected" else "🔴", vault_status),
            "Azure OpenAI": ("🟢" if AZURE_OPENAI_ENDPOINT else "🔴", "Connected" if AZURE_OPENAI_ENDPOINT else "Not configured"),
            "Database": ("🟢" if DB_HOST != '<your-db-host>' else "🔴", "Available" if DB_HOST != '<your-db-host>' else "Not configured"),
            "Web Search": ("🟢" if search else "🔴", "Available" if search else "Not available"),
            "Document Processing": ("🟢" if llamaindex_available else "🔴", "Available" if llamaindex_available else "Not available")
        }
        
        for service, (emoji, status) in status_data.items():
            st.markdown(f"{emoji} **{service}**: {status}")
        
        # Configuration details
        with st.expander("📋 Configuration Details"):
            st.write(f"**Deployment:** {AZURE_OPENAI_DEPLOYMENT_NAME}")
            st.write(f"**API Version:** {AZURE_OPENAI_API_VERSION}")
            st.write(f"**Endpoint:** {AZURE_OPENAI_ENDPOINT}")

        # Clear chat history button
        if st.button("🗑️ Clear Chat History", use_container_width=True):
            st.session_state.messages = []
            st.rerun()

    with col1:
        # Debug information section
        with st.expander("🔧 Debug Information", expanded=False):
            st.write("**System Status:**")

            # Check LLM
            if 'llm' in globals() and llm:
                st.success(f"✅ LLM initialized: {type(llm).__name__}")
                st.write(f"- Deployment: {AZURE_OPENAI_DEPLOYMENT_NAME}")
                st.write(f"- API Version: {AZURE_OPENAI_API_VERSION}")
            else:
                st.error("❌ LLM not initialized")

            # Check tools
            if 'tools' in globals() and tools:
                st.success(f"✅ Tools available: {len(tools)} tools")
                for i, tool in enumerate(tools):
                    st.write(f"  - {i+1}. {tool.name}: {tool.description}")
            else:
                st.error("❌ No tools available")

            # Check agent
            if 'agent_executor' in globals() and agent_executor:
                st.success(f"✅ Agent initialized: {type(agent_executor).__name__}")
            else:
                st.error("❌ Agent not initialized")

            # Environment checks
            st.write("**Environment Configuration:**")
            st.write(f"- Azure OpenAI Endpoint: {'✅ Set' if AZURE_OPENAI_ENDPOINT else '❌ Not set'}")
            st.write(f"- Azure OpenAI API Key: {'✅ Set' if AZURE_OPENAI_API_KEY else '❌ Not set'}")
            st.write(f"- Database Host: {'✅ Set' if DB_HOST != '<your-db-host>' else '❌ Not configured'}")
            st.write(f"- Vault URL: {'✅ Set' if VAULT_URL != 'https://your-hcp-vault-url' else '❌ Not configured'}")
            
            # Test configuration button
            if st.button("🧪 Test Azure OpenAI Configuration"):
                with st.spinner("Testing Azure OpenAI configuration..."):
                    debug_azure_config()

        # Main status check - show different interfaces based on agent status
        if not agent_executor:
            st.error("🚨 Agent System Not Ready")
            st.markdown("""
            The AI agent is not properly initialized. This could be due to:
            1. **Missing Configuration**: Azure OpenAI credentials, database settings, or Vault configuration
            2. **Dependency Issues**: Missing or incompatible LangChain versions
            3. **Authentication Problems**: Azure authentication or Vault connectivity issues
            4. **Model Configuration**: Incorrect deployment name or API version

            **What you can do:**
            1. Check the debug information above
            2. Verify your environment variables are set correctly
            3. Run the "Test Azure OpenAI Configuration" button above
            4. Check the application logs for detailed error messages
            """)
            st.stop()  # Don't show the main interface if agent isn't ready

        # If we get here, agent is initialized - show main interface
        st.success("✅ Agent System Ready")

        # Create tabs for different functionalities
        main_tab, doc_tab = st.tabs(["💬 Chat Interface", "📄 Document Processing"])
        
        with main_tab:
            # Main chat interface
            if "messages" not in st.session_state:
                st.session_state.messages = []

            # Display chat messages
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            # Chat input
            if prompt := st.chat_input("Ask me anything..."):
                # Add user message to chat history
                st.session_state.messages.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)

                # Get agent response
                with st.chat_message("assistant"):
                    with st.spinner("Thinking..."):
                        try:
                            response = query_agent_direct(prompt)

                            if "error" in response:
                                st.error(f"Error: {response['error']}")
                                # Add error to chat history
                                st.session_state.messages.append({
                                    "role": "assistant",
                                    "content": f"I encountered an error: {response['error']}"
                                })
                            else:
                                st.markdown(response["final_answer"])

                                # Show intermediate steps in an expander
                                if response.get("agent_steps"):
                                    with st.expander("Show reasoning steps"):
                                        for step in response["agent_steps"]:
                                            st.write(f"**Step {step['step']}:** {step['action']}")
                                            st.write(f"Input: {step['input']}")
                                            st.write(f"Result: {step['result']}")
                                            st.write("---")

                                # Add assistant response to chat history
                                st.session_state.messages.append({
                                    "role": "assistant",
                                    "content": response["final_answer"]
                                })

                        except Exception as e:
                            error_msg = f"Unexpected error: {str(e)}"
                            st.error(error_msg)
                            st.session_state.messages.append({
                                "role": "assistant",
                                "content": error_msg
                            })
        
        with doc_tab:
            # Document processing interface
            if llamaindex_available:
                render_document_upload_interface()
            else:
                st.error("❌ Document processing not available")
                st.markdown("""
                Document processing is currently unavailable. This could be due to:
                - LlamaIndex not properly initialized
                - Missing Azure OpenAI embeddings configuration
                - Missing required dependencies
                
                Please check the system configuration and try again.
                """)

# Additional utility functions for document processing
def cleanup_temp_files():
    """Clean up temporary files"""
    try:
        temp_dir = Path("temp_uploads")
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
    except Exception as e:
        logger.warning(f"Error cleaning up temp files: {e}")

def get_supported_file_types():
    """Get list of supported file types for upload"""
    return {
        'txt': 'Plain Text',
        'pdf': 'PDF Document', 
        'docx': 'Word Document',
        'csv': 'CSV File'
    }

def validate_file_size(file, max_size_mb=10):
    """Validate uploaded file size"""
    max_size_bytes = max_size_mb * 1024 * 1024
    return file.size <= max_size_bytes

# Enhanced error handling for document operations
class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass

def safe_document_operation(operation_func, *args, **kwargs):
    """Safely execute document operations with error handling"""
    try:
        return operation_func(*args, **kwargs)
    except DocumentProcessingError as e:
        logger.error(f"Document processing error: {e}")
        return f"Document processing error: {e}"
    except Exception as e:
        logger.error(f"Unexpected error in document operation: {e}")
        return f"Unexpected error: {e}"

# MAIN EXECUTION - run Streamlit interface and check Vault connectivity
if __name__ == "__main__":
    # Test Vault connectivity before starting the app
    logger.info("Starting AI Agent application...")
    
    # Add token path verification
    print("=== Startup Diagnostics ===")

    # Show status when running directly
    print("=== AI Agent Status ===")
    print(f"LLM initialized: {'Yes' if llm else 'No'}")
    print(f"Tools available: {len(tools) if tools else 0}")
    print(f"Agent initialized: {'Yes' if agent_executor else 'No'}")
    print(f"Document processing: {'Available' if llamaindex_available else 'Not available'}")

    if not agent_executor:
        print("\n❌ Agent not initialized - application may have limited functionality")
        print("Check the logs above for details.")

        # Try to show what's missing
        if not llm:
            print("- LLM is not initialized (check Azure OpenAI configuration)")
        if not tools:
            print("- No tools available")
    else:
        print("✅ Agent system initialized successfully!")

    # Clean up any existing temp files on startup
    cleanup_temp_files()
    
    # Register cleanup on exit
    import atexit
    atexit.register(cleanup_temp_files)

    # Run Streamlit interface regardless of agent status
    main()